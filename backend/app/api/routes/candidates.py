import csv
import html
import io
import logging
import mimetypes
import re
from datetime import datetime, timedelta, timezone
from pathlib import PurePosixPath
from pathlib import Path
from urllib.error import URLError
from urllib.parse import quote, urlparse, urlunparse
from urllib.request import urlopen
from uuid import uuid4

from docx import Document
from fastapi import APIRouter, Body, Depends, File, HTTPException, Query, Response, UploadFile, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy.orm.attributes import flag_modified

from app.api.deps import get_current_hr_user
from app.core.config import get_settings
from app.core.security import hash_token
from app.db.session import get_db
from app.models.candidate import Candidate, CandidateStatus, ensure_candidate_selection_column
from app.models.hr_user import HRUser
from app.models.interview_answer import InterviewAnswer
from app.models.interview import Interview
from app.models.interview_session import InterviewSession, InterviewSessionStatus
from app.models.job_description import JobDescription
from app.models.resume_score import ResumeScore
from app.schemas.candidate import (
    BulkCandidateDeleteRequest,
    CandidateBulkAction,
    CandidateDetailResponse,
    InterviewTrackingItem,
    CandidateListResponse,
    CandidateRead,
    ScreeningRequest,
    ScreeningResponse,
    CandidateStatusUpdate,
    CandidateUploadResponse,
    ResumeScoreResponse,
    ScreeningResultItem,
    SelectionUpdate,
    SendInterviewRequest,
    SendInterviewResponse,
)
from app.services.email_service import send_interview_email
from app.services.activity_service import log_activity
from app.services.interview_base import generate_base_questions
from app.services.resume_parser import parse_and_extract
from app.services.screening_service import run_screening_for_candidates
from app.services.storage import upload_bytes

router = APIRouter()
logger = logging.getLogger(__name__)
EMAIL_RE = re.compile(r'^[^@\s]+@[^@\s]+\.[^@\s]+$')
BLOCKED_TEST_DOMAINS = {'example.com', 'test.com', 'invalid.com'}


def _get_owned_jd(db: Session, jd_id: int, user_id: int) -> JobDescription:
    jd = db.query(JobDescription).filter(JobDescription.id == jd_id, JobDescription.hr_user_id == user_id).first()
    if jd is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='JD not found')
    return jd


def _get_owned_candidate(db: Session, candidate_id: int, user_id: int) -> Candidate:
    ensure_candidate_selection_column(db)
    candidate = (
        db.query(Candidate)
        .join(JobDescription, Candidate.job_description_id == JobDescription.id)
        .options(joinedload(Candidate.resume_score), joinedload(Candidate.interview_session))
        .filter(Candidate.id == candidate_id, JobDescription.hr_user_id == user_id)
        .first()
    )
    if candidate is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='Candidate not found')
    return candidate


def _is_deliverable_email(email: str) -> tuple[bool, str]:
    value = (email or '').strip().lower()
    if not value:
        return False, 'Email missing'
    if not EMAIL_RE.match(value):
        return False, 'Email format invalid'
    domain = value.split('@', maxsplit=1)[1]
    if domain in BLOCKED_TEST_DOMAINS:
        return False, f"Email domain '{domain}' is not deliverable"
    if value.startswith('unknown+'):
        return False, 'Placeholder email detected'
    return True, ''


def _resume_filename(resume_url: str | None) -> str | None:
    if not resume_url:
        return None
    path = PurePosixPath(resume_url.split('?', maxsplit=1)[0])
    return path.name or None


def _projects_to_text(projects: list[dict] | None) -> str | None:
    if not projects:
        return None
    lines: list[str] = []
    for project in projects:
        title = str(project.get('title') or '').strip()
        description = str(project.get('description') or '').strip()
        if title and description:
            lines.append(f'{title}: {description}')
        elif title:
            lines.append(title)
        elif description:
            lines.append(description)
    return '\n'.join(lines).strip() or None


def _fetch_resume_bytes(resume_url: str) -> bytes:
    parsed = urlparse(resume_url)
    safe_url = urlunparse(parsed._replace(path=quote(parsed.path, safe="/:@!$&'()*+,;=")))
    with urlopen(safe_url) as response:
        return response.read()


def _docx_preview_html(docx_bytes: bytes, title: str) -> str:
    document = Document(io.BytesIO(docx_bytes))
    body: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        escaped = html.escape(text)
        style_name = (paragraph.style.name if paragraph.style else '').lower()

        if style_name.startswith('heading'):
            level_suffix = ''.join(ch for ch in style_name if ch.isdigit())
            level = min(max(int(level_suffix or '2'), 1), 6)
            body.append(f'<h{level}>{escaped}</h{level}>')
        else:
            body.append(f'<p>{escaped}</p>')

    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = ''.join(f'<td>{html.escape(cell.text.strip())}</td>' for cell in row.cells)
            rows.append(f'<tr>{cells}</tr>')
        if rows:
            body.append(f'<table>{"".join(rows)}</table>')

    if not body:
        body.append('<p>No preview content available for this resume.</p>')

    return f"""<!DOCTYPE html>
<html lang="en">
  <head>
    <meta charset="utf-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1" />
    <title>{html.escape(title)}</title>
    <style>
      body {{
        margin: 0;
        padding: 24px;
        font-family: Arial, sans-serif;
        color: #111827;
        background: #ffffff;
        line-height: 1.6;
      }}
      h1, h2, h3, h4, h5, h6 {{
        margin: 0 0 12px;
        line-height: 1.3;
      }}
      p {{
        margin: 0 0 12px;
        white-space: pre-wrap;
      }}
      table {{
        width: 100%;
        border-collapse: collapse;
        margin: 12px 0;
      }}
      td {{
        border: 1px solid #e5e7eb;
        padding: 8px 10px;
        vertical-align: top;
      }}
    </style>
  </head>
  <body>
    {''.join(body)}
  </body>
</html>"""


@router.post('/jd/{jd_id}/candidates/upload', response_model=CandidateUploadResponse)
async def upload_candidates(
    jd_id: int,
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> CandidateUploadResponse:
    jd = _get_owned_jd(db, jd_id, current_user.id)
    if jd.hiring_status != 'active':
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail='Applications are closed for this job. New candidates cannot be uploaded.',
        )
    ensure_candidate_selection_column(db)
    settings = get_settings()

    uploaded = 0
    failed: list[dict[str, str]] = []
    allowed_suffixes = {'.pdf', '.docx'}

    for file in files:
        suffix = Path(file.filename or '').suffix.lower()
        if suffix not in allowed_suffixes:
            failed.append({'file': file.filename or 'unknown', 'error': 'Only PDF and DOCX are supported'})
            continue

        try:
            content = await file.read()
            parsed = parse_and_extract(content)
            if parsed.get('_extraction_failed'):
                logger.warning(
                    f"[PARSER] Extraction failed for {file.filename}: "
                    f"{parsed.get('_failure_reason', 'unknown')}"
                )
            email = (parsed.get('email') or '').strip()
            file_path = f'{jd_id}/{uuid4().hex}_{file.filename}'
            resume_url = upload_bytes(
                settings.supabase_bucket_resumes,
                file_path,
                content,
                content_type=file.content_type or 'application/octet-stream',
            )

            candidate = Candidate(
                job_description_id=jd.id,
                full_name=parsed['full_name'],
                email=email,
                phone=parsed['phone'],
                resume_url=resume_url,
                resume_text=parsed['text'],
                extracted_skills=parsed['skills'],
                extracted_experience_years=parsed['experience_years'],
                extracted_education=parsed['education'],
                extracted_projects=parsed['projects'],
                status=CandidateStatus.UPLOADED,
            )
            db.add(candidate)
            db.commit()
            uploaded += 1
        except Exception as exc:
            db.rollback()
            failed.append({'file': file.filename or 'unknown', 'error': str(exc)})

    if uploaded:
        try:
            log_activity(
                db,
                current_user.id,
                'resumes_uploaded',
                f'{uploaded} resume(s) uploaded to {jd.title}',
            )
        except Exception:
            pass

    return CandidateUploadResponse(uploaded=uploaded, failed=failed)


@router.get('/jd/{jd_id}/candidates', response_model=CandidateListResponse)
def list_candidates(
    jd_id: int,
    min_score: float | None = Query(default=None, ge=0, le=100),
    max_score: float | None = Query(default=None, ge=0, le=100),
    skill: str | None = None,
    min_experience: float | None = Query(default=None, ge=0),
    max_experience: float | None = Query(default=None, ge=0),
    education_level: str | None = None,
    status_filter: CandidateStatus | None = None,
    passed_only: bool = False,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> CandidateListResponse:
    _get_owned_jd(db, jd_id, current_user.id)
    ensure_candidate_selection_column(db)

    query = (
        db.query(Candidate)
        .join(JobDescription, Candidate.job_description_id == JobDescription.id)
        .options(joinedload(Candidate.resume_score))
        .filter(Candidate.job_description_id == jd_id, JobDescription.hr_user_id == current_user.id)
    )

    if status_filter:
        query = query.filter(Candidate.status == status_filter)
    if min_experience is not None:
        query = query.filter(Candidate.extracted_experience_years >= min_experience)
    if max_experience is not None:
        query = query.filter(Candidate.extracted_experience_years <= max_experience)
    if education_level:
        query = query.filter(Candidate.extracted_education == education_level)

    items = []
    for candidate in query.order_by(Candidate.created_at.desc()).all():
        score = candidate.resume_score
        if min_score is not None and (score is None or score.overall_score < min_score):
            continue
        if max_score is not None and (score is None or score.overall_score > max_score):
            continue
        if passed_only and (score is None or not score.passed):
            continue
        if skill:
            if skill.lower() not in {s.lower() for s in candidate.extracted_skills or []}:
                continue
        items.append(candidate)

    return CandidateListResponse(items=[CandidateRead.model_validate(x) for x in items], total=len(items))


@router.post('/jd/{jd_id}/screen', response_model=ScreeningResponse)
def run_screening(
    jd_id: int,
    payload: ScreeningRequest,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> ScreeningResponse:
    jd = _get_owned_jd(db, jd_id, current_user.id)
    if not payload.candidate_ids:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='No candidate IDs provided')
    candidate_ids = list(dict.fromkeys(payload.candidate_ids))
    results = run_screening_for_candidates(
        db=db,
        jd_id=jd_id,
        candidate_ids=candidate_ids,
        skill_importance=payload.skill_importance,
        exp_importance=payload.exp_importance,
        edu_importance=payload.edu_importance,
        project_importance=payload.project_importance,
        active_skills=payload.active_skills,
        threshold=payload.threshold,
    )
    passed = [item for item in results if item['passed']]
    failed = [item for item in results if not item['passed']]

    if results:
        try:
            log_activity(
                db,
                current_user.id,
                'screening_done',
                f'Screened {len(results)} candidate(s) for {jd.title} - {len(passed)} passed, {len(failed)} failed',
            )
        except Exception:
            pass

    return ScreeningResponse(
        results=[ScreeningResultItem(**item) for item in results],
        total_screened=len(results),
        total_passed=len(passed),
        total_failed=len(failed),
    )


@router.get('/candidates/{candidate_id}/scores/{jd_id}', response_model=ResumeScoreResponse)
def get_candidate_scores(
    candidate_id: int,
    jd_id: int,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> ResumeScore:
    candidate = _get_owned_candidate(db, candidate_id, current_user.id)
    if candidate.job_description_id != jd_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='Score not found')

    score = db.query(ResumeScore).filter(ResumeScore.candidate_id == candidate_id, ResumeScore.jd_id == jd_id).first()
    if score is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='Score not found')
    return score


@router.post('/candidates/shortlist', response_model=CandidateListResponse)
def shortlist_candidates(
    payload: CandidateBulkAction,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> CandidateListResponse:
    ensure_candidate_selection_column(db)
    candidates = (
        db.query(Candidate)
        .join(JobDescription, Candidate.job_description_id == JobDescription.id)
        .options(joinedload(Candidate.resume_score))
        .filter(Candidate.id.in_(payload.candidate_ids), JobDescription.hr_user_id == current_user.id)
        .all()
    )
    for candidate in candidates:
        candidate.status = CandidateStatus.SHORTLISTED
    db.commit()
    if candidates:
        try:
            log_activity(
                db,
                current_user.id,
                'shortlisted',
                f'{len(candidates)} candidate(s) shortlisted',
            )
        except Exception:
            pass
    return CandidateListResponse(items=[CandidateRead.model_validate(x) for x in candidates], total=len(candidates))


@router.post('/candidates/send-interviews', response_model=SendInterviewResponse)
def send_interviews(
    payload: SendInterviewRequest,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> SendInterviewResponse:
    ensure_candidate_selection_column(db)
    now = datetime.now(timezone.utc)
    candidate_ids = list(dict.fromkeys(payload.candidate_ids))

    candidates = (
        db.query(Candidate)
        .join(JobDescription, Candidate.job_description_id == JobDescription.id)
        .options(joinedload(Candidate.job_description), joinedload(Candidate.interview_session))
        .filter(Candidate.id.in_(candidate_ids), JobDescription.hr_user_id == current_user.id)
        .all()
    )

    sent = 0
    results: list[dict[str, object]] = []

    for candidate in candidates:
        try:
            ok, reason = _is_deliverable_email(candidate.email)
            if not ok:
                results.append({'candidate_id': candidate.id, 'success': False, 'error': reason})
                continue

            if candidate.status not in {CandidateStatus.SHORTLISTED, CandidateStatus.INTERVIEW_SENT}:
                results.append(
                    {
                        'candidate_id': candidate.id,
                        'success': False,
                        'error': 'Candidate must be shortlisted before sending interview.',
                    }
                )
                continue

            jd = candidate.job_description
            if jd is None:
                results.append({'candidate_id': candidate.id, 'success': False, 'error': 'Job description not found'})
                continue

            existing_invite = (
                db.query(Interview)
                .filter(Interview.candidate_id == candidate.id)
                .order_by(Interview.sent_at.desc())
                .first()
            )
            has_previous_invite = bool(existing_invite and existing_invite.email_sent)
            if has_previous_invite and not payload.force_resend:
                results.append({'candidate_id': candidate.id, 'success': False, 'error': 'already_sent'})
                continue

            questions = generate_base_questions(candidate, jd)
            session = candidate.interview_session
            if session is None:
                session = InterviewSession(candidate_id=candidate.id, status=InterviewSessionStatus.NOT_STARTED, questions=questions)
                db.add(session)
            else:
                session.questions = questions
                session.status = InterviewSessionStatus.NOT_STARTED
                session.video_url = None
                session.started_at = None
                session.completed_at = None
                session.total_score = None
                session.communication_score = None
                session.technical_score = None
                session.behavioral_score = None
                session.confidence_score = None
                session.ai_analysis = None
                session.recommendation = None
                flag_modified(session, 'questions')
                flag_modified(session, 'ai_analysis')
                db.add(session)
                (
                    db.query(InterviewAnswer)
                    .filter(InterviewAnswer.session_id == session.id)
                    .delete(synchronize_session=False)
                )

            logger.info(
                f"[INTERVIEW] Generated {len(questions)} "
                f"questions for candidate {candidate.id}"
            )

            (
                db.query(Interview)
                .filter(Interview.candidate_id == candidate.id, Interview.is_used.is_(False))
                .update({'is_used': True, 'completed_at': datetime.now(timezone.utc)}, synchronize_session=False)
            )

            token = uuid4().hex
            expires_at = now + timedelta(hours=72)
            invite = Interview(
                candidate_id=candidate.id,
                jd_id=jd.id,
                token=token,
                expires_at=expires_at,
                is_used=False,
                force_sent_count=(existing_invite.force_sent_count + 1) if existing_invite and payload.force_resend else 0,
            )
            db.add(invite)

            email_sent = send_interview_email(
                to_email=candidate.email,
                candidate_name=candidate.full_name or candidate.email,
                job_title=jd.title,
                department=jd.department or "",
                interview_token=token,
                expires_at=expires_at,
            )
            invite.email_sent = email_sent

            if email_sent:
                candidate.status = CandidateStatus.INTERVIEW_SENT
                candidate.interview_token = hash_token(token)
                candidate.interview_token_expires = expires_at
                candidate.interview_token_used = False
                db.commit()
                sent += 1
                results.append({'candidate_id': candidate.id, 'success': True, 'token': token})
            else:
                db.commit()
                results.append({'candidate_id': candidate.id, 'success': False, 'error': 'email_failed'})
        except Exception as exc:
            db.rollback()
            results.append({'candidate_id': candidate.id, 'success': False, 'error': str(exc)})

    if sent:
        try:
            log_activity(
                db,
                current_user.id,
                'interviews_sent',
                f'{sent} interview invite(s) sent',
            )
        except Exception:
            pass

    return SendInterviewResponse(sent=sent, results=results)


@router.get('/jd/{jd_id}/interview-results', response_model=list[InterviewTrackingItem])
def get_interview_results(
    jd_id: int,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> list[InterviewTrackingItem]:
    _get_owned_jd(db, jd_id, current_user.id)
    ensure_candidate_selection_column(db)
    now = datetime.now(timezone.utc)

    rows = (
        db.query(Interview, Candidate, ResumeScore)
        .join(Candidate, Interview.candidate_id == Candidate.id)
        .outerjoin(ResumeScore, (ResumeScore.candidate_id == Candidate.id) & (ResumeScore.jd_id == jd_id))
        .join(JobDescription, Candidate.job_description_id == JobDescription.id)
        .filter(Interview.jd_id == jd_id, JobDescription.hr_user_id == current_user.id)
        .order_by(Interview.sent_at.desc())
        .all()
    )

    items: list[InterviewTrackingItem] = []
    for invite, candidate, score in rows:
        status_label = 'Completed' if invite.completed_at else 'Opened' if invite.opened_at else 'Pending'
        items.append(
            InterviewTrackingItem(
                candidate_id=candidate.id,
                candidate_name=candidate.full_name,
                email=candidate.email,
                overall_score=score.overall_score if score else None,
                sent_at=invite.sent_at,
                opened_at=invite.opened_at,
                completed_at=invite.completed_at,
                is_used=invite.is_used,
                status=status_label,
                can_force_resend=bool(invite.is_used or invite.expires_at < now),
            )
        )

    return items


@router.get('/jd/{jd_id}/candidates/export')
def export_candidates(
    jd_id: int,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> StreamingResponse:
    _get_owned_jd(db, jd_id, current_user.id)
    ensure_candidate_selection_column(db)
    candidates = (
        db.query(Candidate)
        .options(joinedload(Candidate.resume_score))
        .filter(Candidate.job_description_id == jd_id)
        .order_by(Candidate.created_at.desc())
        .all()
    )

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Candidate ID', 'Name', 'Email', 'Status', 'Overall Score', 'Passed', 'Experience', 'Education'])
    for c in candidates:
        writer.writerow([
            c.id,
            c.full_name,
            c.email,
            c.status.value,
            c.resume_score.overall_score if c.resume_score else '',
            c.resume_score.passed if c.resume_score else '',
            c.extracted_experience_years,
            c.extracted_education,
        ])

    output.seek(0)
    filename = f'jd_{jd_id}_candidates_{datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")}.csv'
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type='text/csv',
        headers={'Content-Disposition': f'attachment; filename={filename}'},
    )


@router.put('/candidates/{candidate_id}/status', response_model=CandidateRead)
def update_candidate_status(
    candidate_id: int,
    payload: CandidateStatusUpdate,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> Candidate:
    candidate = _get_owned_candidate(db, candidate_id, current_user.id)
    candidate.status = payload.status
    db.commit()
    db.refresh(candidate)
    return candidate


@router.patch('/candidates/{candidate_id}/selection')
def update_candidate_selection(
    candidate_id: int,
    payload: SelectionUpdate,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> dict[str, str | None]:
    candidate = _get_owned_candidate(db, candidate_id, current_user.id)
    candidate.selection = payload.selection
    db.commit()
    return {'selection': candidate.selection}


@router.get('/candidates/{candidate_id}', response_model=CandidateDetailResponse)
def get_candidate(
    candidate_id: int,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> CandidateDetailResponse:
    candidate = _get_owned_candidate(db, candidate_id, current_user.id)
    return CandidateDetailResponse(
        id=candidate.id,
        jd_id=candidate.job_description_id,
        full_name=candidate.full_name,
        email=candidate.email,
        phone=candidate.phone,
        skills=candidate.extracted_skills or [],
        years_of_experience=candidate.extracted_experience_years,
        education_level=candidate.extracted_education,
        projects=_projects_to_text(candidate.extracted_projects),
        resume_url=candidate.resume_url,
        resume_filename=_resume_filename(candidate.resume_url),
        resume_text=candidate.resume_text,
        status=candidate.status.value,
        selection=candidate.selection,
        created_at=candidate.created_at,
        updated_at=candidate.created_at,
        resume_score=ResumeScoreResponse.model_validate(candidate.resume_score) if candidate.resume_score else None,
    )


@router.get('/candidates/{candidate_id}/resume-preview')
def get_candidate_resume_preview(
    candidate_id: int,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> Response:
    candidate = _get_owned_candidate(db, candidate_id, current_user.id)

    if not candidate.resume_url:
      raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='Resume file not available')

    filename = _resume_filename(candidate.resume_url) or 'resume'
    suffix = Path(filename).suffix.lower()

    try:
        content = _fetch_resume_bytes(candidate.resume_url)
    except URLError as exc:
        raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail='Unable to load resume preview') from exc

    if suffix == '.docx':
        preview_html = _docx_preview_html(content, f'{candidate.full_name} resume')
        return Response(content=preview_html.encode('utf-8'), media_type='text/html')

    media_type = 'application/pdf' if suffix == '.pdf' else (mimetypes.guess_type(filename)[0] or 'application/octet-stream')
    return Response(
        content=content,
        media_type=media_type,
        headers={'Content-Disposition': f'inline; filename="{filename}"'},
    )


@router.delete('/candidates/bulk')
def bulk_delete_candidates(
    payload: BulkCandidateDeleteRequest = Body(...),
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> dict[str, int]:
    candidate_ids = list(dict.fromkeys(payload.candidate_ids))
    if not candidate_ids:
        return {'deleted': 0}

    ensure_candidate_selection_column(db)
    candidates = (
        db.query(Candidate)
        .join(JobDescription, Candidate.job_description_id == JobDescription.id)
        .options(joinedload(Candidate.resume_score), joinedload(Candidate.interview_session))
        .filter(Candidate.id.in_(candidate_ids))
        .all()
    )

    if len(candidates) != len(candidate_ids):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='One or more candidates not found')

    if any(candidate.job_description.hr_user_id != current_user.id for candidate in candidates):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail='One or more candidates do not belong to you')

    for candidate in candidates:
        db.delete(candidate)

    db.commit()
    return {'deleted': len(candidates)}


@router.delete('/candidates/{candidate_id}')
def delete_candidate(
    candidate_id: int,
    db: Session = Depends(get_db),
    current_user: HRUser = Depends(get_current_hr_user),
) -> dict[str, str]:
    candidate = _get_owned_candidate(db, candidate_id, current_user.id)
    db.delete(candidate)
    db.commit()
    return {'message': 'Candidate deleted'}
